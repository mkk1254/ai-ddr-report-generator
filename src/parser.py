"""Document parser for PDF, DOCX, and plain text inspection reports."""

import re
from pathlib import Path
from typing import Optional


def parse_pdf(path: Path) -> str:
    """Extract text from a PDF file. Uses pdfplumber for tables, PyPDF2 as fallback."""
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            pages = []
            for page in pdf.pages:
                # Extract text
                text = page.extract_text()
                if text:
                    pages.append(text)
                # Try to extract tables if text is sparse
                tables = page.extract_tables()
                if tables and (not text or len(text.strip()) < 100):
                    for table in tables:
                        for row in table:
                            if row:
                                pages.append(" | ".join(str(cell or "") for cell in row))
            return "\n\n".join(pages) if pages else ""
    except ImportError:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(path)
            return "\n\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            raise RuntimeError(f"Failed to parse PDF: {e}") from e
    except Exception as e:
        raise RuntimeError(f"Failed to parse PDF: {e}") from e


def parse_docx(path: Path) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [str(cell.text or "").strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts) if parts else ""
    except ImportError as e:
        raise RuntimeError("python-docx is required for DOCX files") from e
    except Exception as e:
        raise RuntimeError(f"Failed to parse DOCX: {e}") from e


def parse_txt(path: Path) -> str:
    """Read plain text file."""
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
    last_error = None
    for enc in encodings:
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError as e:
            last_error = e
    raise RuntimeError(f"Could not decode text file with any encoding: {last_error}") from last_error


def parse_document(path: Path) -> str:
    """Parse a document (PDF, DOCX, or TXT) and return normalized text."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = parse_pdf(path)
    elif suffix in (".docx", ".doc"):
        text = parse_docx(path)
    elif suffix in (".txt", ".md", ""):
        text = parse_txt(path)
    else:
        raise ValueError(f"Unsupported file format: {suffix}. Use PDF, DOCX, or TXT.")

    return _normalize_text(text)


def _normalize_text(text: str) -> str:
    """Normalize extracted text: collapse excessive whitespace, preserve structure."""
    if not text or not text.strip():
        return ""
    # Collapse multiple newlines to double newline
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse multiple spaces to single
    text = re.sub(r" {2,}", " ", text)
    return text.strip()
